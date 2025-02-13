import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import os
from tkinter.ttk import Progressbar
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class KIDGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("KID Document Generator")
        self.create_widgets()

    def create_widgets(self):
        tk.Label(self.root, text="1. Select Data Input File (.csv):").pack(anchor='w', padx=10)
        self.data_file_button = tk.Button(self.root, text="Browse", command=self.select_data_file)
        self.data_file_button.pack(padx=10, pady=5, anchor='w')

        self.data_file_label = tk.Label(self.root, text="", anchor="w", fg="blue")
        self.data_file_label.pack(padx=10, pady=5, anchor='w')

        tk.Label(self.root, text="2. Select KID Template:").pack(anchor='w', padx=10)
        self.template_button = tk.Button(self.root, text="Browse", command=self.select_template_file)
        self.template_button.pack(padx=10, pady=5, anchor='w')

        self.template_label = tk.Label(self.root, text="", anchor="w", fg="blue")
        self.template_label.pack(padx=10, pady=5, anchor='w')

        tk.Label(self.root, text="3. Set Save Destination:").pack(anchor='w', padx=10)
        self.save_button = tk.Button(self.root, text="Browse", command=self.set_save_destination)
        self.save_button.pack(padx=10, pady=5, anchor='w')

        self.save_label = tk.Label(self.root, text="", anchor="w", fg="blue")
        self.save_label.pack(padx=10, pady=5, anchor='w')

        self.progress = Progressbar(self.root, orient="horizontal", length=300, mode="determinate")
        self.progress.pack(pady=10)

        self.generate_button = tk.Button(self.root, text="Generate KID Documents", command=self.generate_documents, state=tk.DISABLED)
        self.generate_button.pack(pady=10)

        self.data_file_path = ""
        self.template_file_path = ""
        self.save_directory = ""

    def select_data_file(self):
        self.data_file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        self.update_label(self.data_file_label, self.data_file_path)

    def select_template_file(self):
        self.template_file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        self.update_label(self.template_label, self.template_file_path)

    def set_save_destination(self):
        self.save_directory = filedialog.askdirectory()
        self.update_label(self.save_label, self.save_directory)

    def update_label(self, label, path):
        if path:
            label.config(text=f"Selected: {path}", fg="green")  # Update the label to show the path and change the text color to green
        else:
            label.config(text="", fg="blue")  # Reset the label if no path is selected
        self.check_ready_to_generate()

    def check_ready_to_generate(self):
        if self.data_file_path and self.template_file_path and self.save_directory:
            self.generate_button.config(state=tk.NORMAL)

    def replace_placeholders(self, doc, replacements):
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    # Replace placeholder text
                    paragraph.text = paragraph.text.replace(placeholder, value)

                    # Set font properties and alignment
                    for run in paragraph.runs:
                        run.font.name = 'Helvetica'
                        run.font.size = Pt(10)  # Set font size to 10 pt
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the paragraph

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            # Replace placeholder text
                            cell.text = cell.text.replace(placeholder, value)

                            # Set font properties and alignment in the cell
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Helvetica'
                                    run.font.size = Pt(10)  # Set font size to 10 pt
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the paragraph

    def calculate_ni(self, monthly_rate, ni_lower_threshold=542, ni_primary_threshold=1048, ni_upper_threshold=4189,
                     ni_rate_13=0.138, ni_rate_2=0.02):
        """
        Calculate NI contributions based on earnings above the threshold.
        :param monthly_rate: The gross monthly rate
        :param ni_lower_threshold: The lower earnings limit for the month
        :param ni_primary_threshold: The primary threshold for employee contributions
        :param ni_upper_threshold: The upper earnings limit for employee contributions
        :param ni_rate_13: The 13.8% NI rate applied between the secondary and upper earnings limits
        :param ni_rate_2: The 2% NI rate applied above the upper earnings limit
        :return: The calculated NI deduction
        """
        ni_deduction = 0.0

        # Earnings between lower and primary threshold taxed at 13.8% (employer contribution)
        if monthly_rate > ni_lower_threshold:
            ni_taxable = min(monthly_rate, ni_primary_threshold) - ni_lower_threshold
            ni_deduction += ni_taxable * ni_rate_13

        # Earnings between primary and upper threshold taxed at 13.8% (employee contribution)
        if monthly_rate > ni_primary_threshold:
            ni_taxable = min(monthly_rate, ni_upper_threshold) - ni_primary_threshold
            ni_deduction += ni_taxable * ni_rate_13

        # Earnings above the upper threshold taxed at 2% (employee contribution)
        if monthly_rate > ni_upper_threshold:
            ni_taxable = monthly_rate - ni_upper_threshold
            ni_deduction += ni_taxable * ni_rate_2

        return ni_deduction

    def calculate_tax(self, daily_rate, working_days, tax_thresholds=None, tax_rates=None):
        """
        Calculate tax deduction based on UK tax bands.
        :param daily_rate: The daily pay rate
        :param working_days: Number of working days in the month
        :param tax_thresholds: Dictionary of tax thresholds
        :param tax_rates: Dictionary of tax rates for each band
        :return: The total monthly tax deduction
        """
        # Set default thresholds and rates if not provided
        if tax_thresholds is None:
            tax_thresholds = {
                "personal_allowance": 12570,  # Personal allowance in GBP (annually)
                "basic_rate": 37700,  # Basic rate threshold (annually)
                "higher_rate": 125140,  # Higher rate threshold (annually)
                "additional_rate": 125140  # Additional rate threshold (annually)
            }
        if tax_rates is None:
            tax_rates = {
                "personal_allowance": 0.0,  # 0% tax
                "basic_rate": 0.2,  # 20% tax
                "higher_rate": 0.4,  # 40% tax
                "additional_rate": 0.45  # 45% tax
            }

        monthly_rate = daily_rate * working_days  # Calculate monthly earnings
        annual_income = monthly_rate * 12  # Convert to annual income
        print(f"Monthly Rate: {monthly_rate}")
        print(f"Annual Income: {annual_income}")  # Debugging the annual income

        # Initialize tax deduction
        tax_deduction = 0.0

        # Apply tax bands based on annual income
        if annual_income <= tax_thresholds["personal_allowance"]:
            print(f"Annual income within personal allowance, no tax.")  # Debug
            tax_deduction = 0  # No tax for income within personal allowance
        else:
            taxable_income = annual_income - tax_thresholds["personal_allowance"]
            print(f"Taxable Income: {taxable_income}")  # Debug

            # Calculate tax progressively based on the thresholds and rates
            if taxable_income <= tax_thresholds["basic_rate"]:
                tax_deduction = taxable_income * tax_rates["basic_rate"]
            else:
                # Income between personal allowance and basic rate threshold
                tax_deduction = (tax_thresholds["basic_rate"] - tax_thresholds["personal_allowance"]) * tax_rates[
                    "basic_rate"]

                # Income between basic rate and higher rate threshold
                if taxable_income <= tax_thresholds["higher_rate"]:
                    tax_deduction += (taxable_income - tax_thresholds["basic_rate"]) * tax_rates["higher_rate"]
                else:
                    # Income between higher rate and additional rate threshold
                    tax_deduction += (tax_thresholds["higher_rate"] - tax_thresholds["basic_rate"]) * tax_rates[
                        "higher_rate"]

                    # Income above the higher rate threshold
                    tax_deduction += (taxable_income - tax_thresholds["higher_rate"]) * tax_rates["additional_rate"]

        # Calculate monthly tax deduction based on the annual deduction
        monthly_tax_deduction = tax_deduction / 12  # Divide by 12 to get monthly tax
        print(f"Monthly Tax Deduction: {monthly_tax_deduction}")  # Debug

        return monthly_tax_deduction

    def calculate_pension_contribution(self, daily_rate, working_days, avg_day_pension_percent=0.04281):
        """
        Calculate the monthly pension contribution based on the daily rate and pension percentage.
        :param daily_rate: The daily pay rate
        :param working_days: Number of working days in the month
        :param avg_day_pension_percent: The pension percentage to apply to the daily rate
        :return: The monthly pension contribution
        """
        day_pension_contribution = daily_rate * avg_day_pension_percent
        month_pension_contribution = day_pension_contribution * working_days
        return month_pension_contribution

    def calculate_employer_ni(self, monthly_rate, ni_threshold=758.33, employer_ni_rate=0.138):
        """
        Calculate the employer's National Insurance (NIC) contributions based on monthly earnings.
        :param monthly_rate: The monthly gross earnings
        :param ni_threshold: The threshold above which employer's NIC is applicable
        :param employer_ni_rate: The NIC rate applied above the threshold
        :return: The employer's NIC contribution
        """
        earnings_above_threshold = max(0, monthly_rate - ni_threshold)
        employer_ni_deduction = earnings_above_threshold * employer_ni_rate
        return employer_ni_deduction

    def calculate_employer_pension(self, monthly_rate, pension_threshold=520, employer_pension_rate=0.03):
        """
        Calculate the employer's pension contribution based on qualifying earnings.
        :param monthly_rate: The monthly gross earnings
        :param pension_threshold: The threshold above which pension contributions apply
        :param employer_pension_rate: The employer's pension contribution rate
        :return: The employer's pension contribution
        """
        qualifying_earnings = max(0, monthly_rate - pension_threshold)
        employer_pension_contribution = qualifying_earnings * employer_pension_rate
        return employer_pension_contribution



    def generate_documents(self):
        global working_days
        working_days = 20
        umbrella_fee = 20.00  # Default umbrella fee

        # Read the CSV file and skip the second row
        data = pd.read_csv(self.data_file_path, skiprows=[1])
        total_rows = len(data)
        successful_generations = 0


        for index, row in data.iterrows():
            # Skip Row 2
            # (No need for manual check here since it's already handled by skiprows)

            # Retrieve data from the row
            cand_full_name = row.get("Candidate", "")
            assignment_id = row.get("ID", "")  # Assuming 'AssignmentID' exists
            contract_start_date = row.get("Start Date", "")
            contract_end_date = row.get("End Date", "")
            daily_rate = float(row.get("Pay Rate", 0.0))
            umbrella_name = row.get("Umbrella Company", "")
            job_title = row.get("Vacancy", "")
            pay_freq = row.get("Pay Unit")

            if pay_freq.lower() == "per month":
                daily_rate = daily_rate / working_days

            if not all([cand_full_name, assignment_id, contract_start_date, contract_end_date]):
                print(f"Skipping row {index + 1} due to missing data")
                continue  # Skip rows with missing critical data

            # Financial calculations
            avg_day_tax_percent = 0.1389
            avg_day_pension_percent = 0.04281
            monthly_rate = daily_rate * working_days  # Monthly rate calculated from daily rate

            # Deductions based on daily rate
            month_nic_deduction = self.calculate_ni(monthly_rate)
            month_tax_deduction = self.calculate_tax(daily_rate, working_days)
            month_pension_contribution = self.calculate_pension_contribution(daily_rate, working_days)

            # Total deductions for the month
            total_deductions = month_tax_deduction + month_nic_deduction + month_pension_contribution + (
                    umbrella_fee * 4)



            # Net pay calculation based on monthly rate
            net_pay = monthly_rate - total_deductions

            # Annual figures
            annual_rate = monthly_rate * 12
            apprenticeship_levy_constant = 0.05
            apprenticeship_levy = daily_rate * apprenticeship_levy_constant

            # Employer's NIC and Pension calculations
            employer_ni_deduction = self.calculate_employer_ni(monthly_rate)
            employer_pension_contribution = self.calculate_employer_pension(monthly_rate)

            # Example rate of pay to you (contractor)
            eg_monthly_rate_to_contractor = monthly_rate - employer_ni_deduction - employer_pension_contribution - apprenticeship_levy

            doc = Document(self.template_file_path)

            # Prepare replacements
            replacements = {
                "{CandidateName}": cand_full_name,
                "{ContractStartDate}": contract_start_date,
                "{ContractEndDate}": contract_end_date,
                "{DailyRate}": f"£{daily_rate:.2f}",
                "{TaxDeduction}": f"£{month_tax_deduction:.2f}",
                "{NICDeduction}": f"£{month_nic_deduction:.2f}",
                "{UmbrellaFee}": f"£{umbrella_fee:.2f} per week",
                "{NetPay}": f"£{net_pay:.2f}",
                "{MonthlyRate}": f"£{monthly_rate:.2f}",
                "{TotalDeductions}": f"£{total_deductions:.2f}",
                "{AssignmentID}": f"{assignment_id}",
                "{UmbrellaName}": f"{umbrella_name}",
                "{JobTitle}": f"{job_title}",
                "{MinWage}": "£11.44 per hour",
                "{PayFreq}": f"{pay_freq}",
                "{ApprenLevy}": f"£{apprenticeship_levy:.2f}",
                "{PensionCont}": f"£{month_pension_contribution:.2f}",
                "{WorkDays}": f"{working_days}",
                "{EmployerNIC}": f"£{employer_ni_deduction:.2f}",
                "{EmployerPension}": f"£{employer_pension_contribution:.2f}",
                "{EgPayToContractor}": f"£{eg_monthly_rate_to_contractor:.2f}"
            }

            # Replace placeholders
            self.replace_placeholders(doc, replacements)

            # Replace / with - in file name
            contract_start_date = contract_start_date.replace("/", "-")
            contract_end_date = contract_end_date.replace("/", "-")

            # Updated file name format: CandidateName_-_AssignmentID_-_StartDate_-_EndDate_-_KIDFORM
            file_name = f"{cand_full_name}_{assignment_id}_{contract_start_date}_{contract_end_date}_KIDFORM.docx"
            save_path = os.path.join(self.save_directory, file_name)

            try:
                # Save the Word document
                doc.save(save_path)
                successful_generations += 1
            except Exception as e:
                print(f"Error generating {file_name}: {e}")

            # Update progress
            self.progress['value'] = ((index + 1) / total_rows) * 100
            self.root.update_idletasks()

        messagebox.showinfo("Generation Complete",
                            f"Successfully generated {successful_generations} of {total_rows} documents.\n"
                            f"Thank you for using this Kid Form Generator\n"
                            f"Kid Form Generator v1.5 - created by Peter Grint - Feb 2025")


if __name__ == "__main__":
    root = tk.Tk()
    app = KIDGeneratorApp(root)
    root.mainloop()
